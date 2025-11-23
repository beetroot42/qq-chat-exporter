#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
提取QQ聊天记录中的RPG会话，并生成彩色Word文档

使用方法：
    python extract_rpg_sessions.py <输入文件> [输出文件]

示例：
    python extract_rpg_sessions.py chat_log.txt output.docx
    python extract_rpg_sessions.py chat_log.txt
"""

import re
import sys
from docx import Document
from docx.shared import RGBColor

# 网名到角色名的映射
NAME_MAPPING = {
    # 网名映射
    '失语': '【神谕圣咏】吾即五声',
    'heavy🐜': '【调停者】阿德勒',
    'Ga1axian': '【九号球】玖渚巡',
    '随波逐流制作委员会': '【黑印】加尔文',
    '梦之海': '【菟丝子】十七',
    '无糖常温百事FES': '【时代】艾帕克',
    # QQ号映射（有时会显示QQ号而非网名）
    '907564155': '【神谕圣咏】吾即五声',
    '2704587599': '【调停者】阿德勒',
    '1456846090': '【九号球】玖渚巡',
    '1695539040': '【黑印】加尔文',
    '1214581195': '【菟丝子】十七',
    '651464169': '【时代】艾帕克'
}

# 角色颜色映射 (RGB格式)
COLOR_MAPPING = {
    '【神谕圣咏】吾即五声': RGBColor(255, 0, 0),      # 红色
    '【调停者】阿德勒': RGBColor(128, 0, 128),        # 紫色
    '【九号球】玖渚巡': RGBColor(139, 69, 19),        # 棕色
    '【黑印】加尔文': RGBColor(0, 0, 0),              # 黑色
    '【菟丝子】十七': RGBColor(0, 0, 255),            # 蓝色
    '【时代】艾帕克': RGBColor(128, 128, 128)         # 灰色
}

def parse_chat_log(file_path):
    """解析聊天记录文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    messages = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # 检测是否是用户名行（以 ':' 结尾）
        if line and line.endswith(':'):
            username = line[:-1]  # 去掉冒号
            timestamp = None
            content = None

            # 读取时间戳
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('时间:'):
                timestamp = lines[i + 1].strip().replace('时间:', '').strip()
                i += 1

            # 读取内容（可能是多行，可能包含空行）
            if i + 1 < len(lines) and lines[i + 1].strip().startswith('内容:'):
                content_lines = [lines[i + 1].strip().replace('内容:', '').strip()]
                i += 1

                # 继续读取后续的内容行
                consecutive_empty = 0
                while i + 1 < len(lines):
                    next_line = lines[i + 1].strip()

                    # 如果是提及行或新消息，停止读取
                    if next_line.startswith('提及:') or (next_line.endswith(':') and i + 2 < len(lines) and lines[i + 2].strip().startswith('时间:')):
                        break

                    # 如果是空行
                    if not next_line:
                        consecutive_empty += 1
                        # 连续2个空行，说明消息结束
                        if consecutive_empty >= 2:
                            break
                        # 保留这个空行，继续读取
                        content_lines.append('')
                        i += 1
                    else:
                        # 非空行，重置连续空行计数
                        consecutive_empty = 0
                        content_lines.append(next_line)
                        i += 1

                # 去除末尾的空行
                while content_lines and not content_lines[-1]:
                    content_lines.pop()

                content = '\n'.join(content_lines)

            # 跳过提及行和其他元数据
            while i + 1 < len(lines) and lines[i + 1].strip() and lines[i + 1].strip().startswith('提及:'):
                i += 1

            if content:
                messages.append({
                    'username': username,
                    'timestamp': timestamp,
                    'content': content
                })

        i += 1

    return messages

def extract_rpg_sessions(messages):
    """提取RPG会话"""
    sessions = []
    in_session = False
    current_session = []
    # 匹配所有格式：
    # ——CST5016/08/14——
    # ——CST5016/08/14/1900时——
    # ——CST5016/10/21/1200——
    # ——CST5017/01/10——
    # ——CST5017/01/11，AST0500——
    start_pattern = re.compile(r'——CST501[67]/\d+/\d+(/\d+时?)?(，[^—]+)?——')

    for msg in messages:
        content = msg['content']

        # 检测开始标记
        if start_pattern.search(content):
            # 如果之前有未完成的会话，先保存它（即使没有save标记）
            if in_session and current_session:
                sessions.append(current_session)
            in_session = True
            current_session = [msg]
        elif in_session:
            current_session.append(msg)
            # 检测结束标记
            if content == '——save——':
                sessions.append(current_session)
                current_session = []
                in_session = False

    # 处理最后一个未完成的会话（如果文件结束时还在会话中）
    if in_session and current_session:
        sessions.append(current_session)

    return sessions

def process_rpg_session(session):
    """处理RPG会话：替换网名，处理时间戳"""
    processed = []
    start_pattern = re.compile(r'——CST501[67]/\d+/\d+(/\d+时?)?(，[^—]+)?——')

    for msg in session:
        # 替换网名为角色名
        username = msg['username']
        if username in NAME_MAPPING:
            username = NAME_MAPPING[username]

        content = msg['content']

        # 判断是否是开始标记或结束标记
        is_start_marker = start_pattern.search(content)
        is_end_marker = content == '——save——'

        # 如果是开始/结束标记，保留时间戳；否则删除
        if is_start_marker or is_end_marker:
            timestamp = msg['timestamp']
        else:
            timestamp = None

        processed.append({
            'username': username,
            'timestamp': timestamp,
            'content': content
        })

    return processed

def create_word_document(sessions, output_path):
    """创建带颜色的Word文档"""
    doc = Document()
    doc.add_heading('RPG会话记录', 0)

    for idx, session in enumerate(sessions, 1):
        # 添加会话标题
        doc.add_heading(f'会话 {idx}', level=1)

        # 处理会话内容
        processed_session = process_rpg_session(session)

        for msg in processed_session:
            # 创建段落
            paragraph = doc.add_paragraph()

            # 添加用户名
            run_username = paragraph.add_run(msg['username'])
            if msg['username'] in COLOR_MAPPING:
                run_username.font.color.rgb = COLOR_MAPPING[msg['username']]
            run_username.bold = True

            # 添加时间戳（如果有）
            if msg['timestamp']:
                run_time = paragraph.add_run(f" [{msg['timestamp']}]")
                if msg['username'] in COLOR_MAPPING:
                    run_time.font.color.rgb = COLOR_MAPPING[msg['username']]

            # 添加冒号
            run_colon = paragraph.add_run(': ')
            if msg['username'] in COLOR_MAPPING:
                run_colon.font.color.rgb = COLOR_MAPPING[msg['username']]

            # 添加内容
            run_content = paragraph.add_run(msg['content'])
            if msg['username'] in COLOR_MAPPING:
                run_content.font.color.rgb = COLOR_MAPPING[msg['username']]

    doc.save(output_path)
    print(f'Word文档已生成：{output_path}')

def main():
    # 解析命令行参数
    if len(sys.argv) < 2:
        print(__doc__)
        print('错误：请提供输入文件路径')
        print('使用方法: python extract_rpg_sessions.py <输入文件> [输出文件]')
        sys.exit(1)

    input_file = sys.argv[1]

    # 如果提供了输出文件名，使用它；否则根据输入文件名自动生成
    if len(sys.argv) >= 3:
        output_file = sys.argv[2]
    else:
        # 自动生成输出文件名：input.txt -> input_rpg.docx
        import os
        base_name = os.path.splitext(input_file)[0]
        output_file = f'{base_name}_rpg.docx'

    print(f'输入文件: {input_file}')
    print(f'输出文件: {output_file}')
    print('-' * 50)

    try:
        print('正在解析聊天记录...')
        messages = parse_chat_log(input_file)
        print(f'共解析到 {len(messages)} 条消息')

        print('正在提取RPG会话...')
        sessions = extract_rpg_sessions(messages)
        print(f'共提取到 {len(sessions)} 个RPG会话')

        if sessions:
            print('正在生成Word文档...')
            create_word_document(sessions, output_file)
            print('✅ 完成！')
        else:
            print('⚠️  未找到RPG会话')
    except FileNotFoundError:
        print(f'❌ 错误：找不到文件 {input_file}')
        sys.exit(1)
    except Exception as e:
        print(f'❌ 处理过程中出错：{e}')
        sys.exit(1)

if __name__ == '__main__':
    main()
